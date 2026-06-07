import json
import math
from pathlib import Path

import matplotlib as mpl
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DATA_PATH = ROOT / "docs" / "latest_run_results.json"
FIG_DIR = ROOT / "docs" / "paper_figures"
DOC_DIR = ROOT / "output" / "doc"
DOC_PATH = DOC_DIR / "Lab6_卷积算子优化实验报告.docx"

METHODS = ["Baseline_FP32", "Optim_FP32", "Optim_FP16", "Optim_WMMA"]
METHOD_LABELS = {
    "Baseline_FP32": "Baseline FP32",
    "Optim_FP32": "Optim FP32",
    "Optim_FP16": "Optim FP16",
    "Optim_WMMA": "Optim WMMA",
}
METHOD_SHORT = {
    "Baseline_FP32": "Baseline",
    "Optim_FP32": "FP32",
    "Optim_FP16": "FP16",
    "Optim_WMMA": "WMMA",
}
COLORS = {
    "Baseline_FP32": "#5B6670",
    "Optim_FP32": "#2A6FBB",
    "Optim_FP16": "#D9822B",
    "Optim_WMMA": "#2F9E7E",
    "accent": "#B23A48",
    "dark": "#1F2933",
    "muted": "#A7B0B8",
    "pale_blue": "#E8F1FA",
    "pale_orange": "#FCEEDB",
    "pale_green": "#E7F5EF",
    "pale_red": "#F9E7EA",
}


def configure_matplotlib():
    mpl.rcParams.update({
        "font.family": "sans-serif",
        "font.sans-serif": [
            "Microsoft YaHei",
            "SimHei",
            "Arial",
            "DejaVu Sans",
            "sans-serif",
        ],
        "axes.unicode_minus": False,
        "svg.fonttype": "none",
        "pdf.fonttype": 42,
        "font.size": 9,
        "axes.spines.top": False,
        "axes.spines.right": False,
        "axes.linewidth": 0.9,
        "legend.frameon": False,
        "figure.facecolor": "white",
        "axes.facecolor": "white",
        "savefig.facecolor": "white",
    })


def load_results():
    return json.loads(DATA_PATH.read_text(encoding="utf-8"))


def save_figure(fig, stem):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    base = FIG_DIR / stem
    for ext in ("png", "svg", "pdf"):
        fig.savefig(base.with_suffix(f".{ext}"), dpi=320, bbox_inches="tight")
    plt.close(fig)
    return base.with_suffix(".png")


def add_panel_label(ax, label):
    ax.text(
        -0.08,
        1.05,
        label,
        transform=ax.transAxes,
        fontsize=12,
        fontweight="bold",
        color=COLORS["dark"],
        va="top",
        ha="left",
    )


def draw_rounded_box(ax, xy, width, height, text, fc, ec, fontsize=10, weight="bold"):
    from matplotlib.patches import FancyBboxPatch

    patch = FancyBboxPatch(
        xy,
        width,
        height,
        boxstyle="round,pad=0.02,rounding_size=0.05",
        linewidth=1.2,
        edgecolor=ec,
        facecolor=fc,
    )
    ax.add_patch(patch)
    ax.text(
        xy[0] + width / 2,
        xy[1] + height / 2,
        text,
        ha="center",
        va="center",
        fontsize=fontsize,
        fontweight=weight,
        color=COLORS["dark"],
        linespacing=1.25,
    )
    return patch


def arrow(ax, start, end, color=None, lw=1.7):
    ax.annotate(
        "",
        xy=end,
        xytext=start,
        arrowprops=dict(
            arrowstyle="-|>",
            lw=lw,
            color=color or COLORS["dark"],
            shrinkA=8,
            shrinkB=8,
            mutation_scale=14,
        ),
    )


def fig_optimization_roadmap():
    fig, ax = plt.subplots(figsize=(9.2, 3.1))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 3)
    ax.axis("off")

    steps = [
        ("直接卷积\nFP32", COLORS["pale_blue"], COLORS["Baseline_FP32"]),
        ("Implicit-GEMM\nFP32", "#EAF3FF", COLORS["Optim_FP32"]),
        ("Implicit-GEMM\nFP16", COLORS["pale_orange"], COLORS["Optim_FP16"]),
        ("WMMA / Tensor Core\nFP16 前向", COLORS["pale_green"], COLORS["Optim_WMMA"]),
    ]
    xs = [0.35, 2.85, 5.35, 7.85]
    for i, (text, fc, ec) in enumerate(steps):
        draw_rounded_box(ax, (xs[i], 1.15), 1.75, 0.9, text, fc, ec, fontsize=10)
        if i < len(steps) - 1:
            arrow(ax, (xs[i] + 1.75, 1.6), (xs[i + 1], 1.6), color=COLORS["muted"])

    labels = [
        "参考实现\n访存无复用",
        "共享内存 + 寄存器 tile\n提升访存局部性",
        "数据位宽减半\n降低显存带宽压力",
        "Warp 级矩阵乘累加\n验证 Tensor Core 路径",
    ]
    for x, label in zip(xs, labels):
        ax.text(x + 0.875, 0.42, label, ha="center", va="center", fontsize=8.5, color="#52616B")

    ax.text(
        0.1,
        2.68,
        "卷积算子优化路线：从朴素卷积到矩阵化、半精度与 Tensor Core",
        fontsize=12,
        fontweight="bold",
        color=COLORS["dark"],
    )
    return save_figure(fig, "fig1_optimization_roadmap")


def fig_extension_workflow():
    fig, ax = plt.subplots(figsize=(9.2, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")

    boxes = [
        ((0.4, 3.25), "Python 模型层\nResNet-18 / Conv2d wrapper", COLORS["pale_blue"], COLORS["Optim_FP32"]),
        ((3.0, 3.25), "PyTorch Extension\nCUDAExtension + pybind11", "#F3F6F8", "#6B7785"),
        ((5.7, 3.25), "C++ 调度层\n参数封装 / Tensor 检查", COLORS["pale_orange"], COLORS["Optim_FP16"]),
        ((8.1, 3.25), "CUDA Kernel\nForward / Backward", COLORS["pale_green"], COLORS["Optim_WMMA"]),
    ]
    for xy, text, fc, ec in boxes:
        draw_rounded_box(ax, xy, 1.65, 0.95, text, fc, ec, fontsize=8.8)

    for x in [2.05, 4.65, 7.35]:
        arrow(ax, (x, 3.72), (x + 0.8, 3.72), color="#7B8794")

    bottom_boxes = [
        ((1.0, 1.1), "输入 Tensor\nNCHW / CUDA", "#FFFFFF", COLORS["muted"]),
        ((3.1, 1.1), "卷积参数\nstride / padding / shape", "#FFFFFF", COLORS["muted"]),
        ((5.45, 1.1), "输出 Tensor\nN×K×Oh×Ow", "#FFFFFF", COLORS["muted"]),
        ((7.65, 1.1), "实验指标\nAccuracy / Throughput", "#FFFFFF", COLORS["muted"]),
    ]
    for xy, text, fc, ec in bottom_boxes:
        draw_rounded_box(ax, xy, 1.75, 0.78, text, fc, ec, fontsize=8.5, weight="normal")

    for start, end in [
        ((1.82, 1.88), (1.15, 3.25)),
        ((3.98, 1.88), (3.82, 3.25)),
        ((6.32, 1.88), (8.65, 3.25)),
        ((8.52, 3.25), (8.52, 1.88)),
    ]:
        arrow(ax, start, end, color="#B0B8C0", lw=1.2)

    ax.text(0.1, 4.68, "PyTorch 自定义 CUDA 卷积算子调用链", fontsize=12, fontweight="bold", color=COLORS["dark"])
    ax.text(
        0.1,
        0.28,
        "该流程将 Python 模型中的 Conv2d 调用映射到 C++/CUDA 扩展模块，使 ResNet-18 推理可以直接调用自定义卷积 kernel。",
        fontsize=8.7,
        color="#52616B",
    )
    return save_figure(fig, "fig2_extension_workflow")


def fig_algorithm_mapping():
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.2), gridspec_kw={"width_ratios": [1.0, 1.1]})
    ax0, ax1 = axes

    for ax in axes:
        ax.set_xticks([])
        ax.set_yticks([])
        ax.set_frame_on(False)

    ax0.set_xlim(0, 6)
    ax0.set_ylim(0, 5)
    add_panel_label(ax0, "a")
    ax0.text(0.1, 4.75, "直接卷积", fontsize=11, fontweight="bold", color=COLORS["dark"])
    draw_rounded_box(ax0, (0.3, 3.1), 1.45, 0.72, "Input\nN,C,H,W", COLORS["pale_blue"], COLORS["Optim_FP32"], fontsize=8.5)
    draw_rounded_box(ax0, (2.25, 3.1), 1.45, 0.72, "Weight\nK,C,R,S", COLORS["pale_orange"], COLORS["Optim_FP16"], fontsize=8.5)
    draw_rounded_box(ax0, (4.2, 3.1), 1.45, 0.72, "Output\nN,K,Oh,Ow", COLORS["pale_green"], COLORS["Optim_WMMA"], fontsize=8.5)
    arrow(ax0, (1.75, 3.45), (2.25, 3.45), color="#8B95A1")
    arrow(ax0, (3.7, 3.45), (4.2, 3.45), color="#8B95A1")

    loop_text = "for n,k,oh,ow:\n  sum = 0\n  for c,r,s:\n    sum += x * w"
    draw_rounded_box(ax0, (1.2, 1.15), 3.65, 1.15, loop_text, "#FFFFFF", "#8B95A1", fontsize=9, weight="normal")
    ax0.text(0.6, 0.45, "问题：访存不规则、输入与权重复用不足、计算访存比低", fontsize=8.5, color="#52616B")

    ax1.set_xlim(0, 7)
    ax1.set_ylim(0, 5)
    add_panel_label(ax1, "b")
    ax1.text(0.1, 4.75, "Implicit-GEMM 映射", fontsize=11, fontweight="bold", color=COLORS["dark"])
    draw_rounded_box(ax1, (0.35, 3.05), 1.3, 0.75, "Weight\nK × CRS", COLORS["pale_orange"], COLORS["Optim_FP16"], fontsize=8.5)
    draw_rounded_box(ax1, (2.1, 3.05), 1.55, 0.75, "Input\nCRS × N·Oh·Ow", COLORS["pale_blue"], COLORS["Optim_FP32"], fontsize=8.5)
    draw_rounded_box(ax1, (4.25, 3.05), 1.55, 0.75, "Output\nK × N·Oh·Ow", COLORS["pale_green"], COLORS["Optim_WMMA"], fontsize=8.5)
    ax1.text(1.82, 3.42, "×", fontsize=15, fontweight="bold", color=COLORS["dark"], ha="center")
    ax1.text(3.93, 3.42, "=", fontsize=15, fontweight="bold", color=COLORS["dark"], ha="center")

    for i, (x, color) in enumerate([(0.45, COLORS["Optim_FP16"]), (2.2, COLORS["Optim_FP32"]), (4.35, COLORS["Optim_WMMA"])]):
        for r in range(4):
            for c in range(5):
                alpha = 0.75 if (r + c + i) % 2 == 0 else 0.35
                rect = plt.Rectangle((x + c * 0.19, 1.15 + r * 0.19), 0.15, 0.15, color=color, alpha=alpha)
                ax1.add_patch(rect)
    ax1.text(0.35, 0.55, "核心：不显式 im2col，而是在 kernel 中计算 CRS 到 NCHW 的地址映射", fontsize=8.5, color="#52616B")
    return save_figure(fig, "fig3_algorithm_mapping")


def fig_wmma_schematic():
    fig, ax = plt.subplots(figsize=(9.2, 4.2))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")

    ax.text(0.1, 4.65, "WMMA / Tensor Core 执行示意", fontsize=12, fontweight="bold", color=COLORS["dark"])
    draw_rounded_box(ax, (0.45, 3.0), 1.65, 0.85, "A fragment\n16×16 FP16", COLORS["pale_orange"], COLORS["Optim_FP16"], fontsize=9)
    draw_rounded_box(ax, (3.0, 3.0), 1.65, 0.85, "B fragment\n16×16 FP16", COLORS["pale_blue"], COLORS["Optim_FP32"], fontsize=9)
    draw_rounded_box(ax, (5.55, 3.0), 1.65, 0.85, "mma_sync\nTensor Core", COLORS["pale_green"], COLORS["Optim_WMMA"], fontsize=9)
    draw_rounded_box(ax, (8.0, 3.0), 1.65, 0.85, "C fragment\n16×16 FP32", "#F3F6F8", "#6B7785", fontsize=9)
    for x in [2.1, 4.65, 7.2]:
        arrow(ax, (x, 3.42), (x + 0.75, 3.42), color="#7B8794")

    draw_rounded_box(ax, (0.75, 1.15), 2.2, 0.9, "1 个 Warp\n32 线程协同", "#FFFFFF", "#8B95A1", fontsize=9, weight="normal")
    draw_rounded_box(ax, (3.95, 1.15), 2.2, 0.9, "每次 16×16×16\n矩阵乘累加", "#FFFFFF", "#8B95A1", fontsize=9, weight="normal")
    draw_rounded_box(ax, (7.05, 1.15), 2.2, 0.9, "大矩阵上显著降低\nFP16 GEMM 时间", "#FFFFFF", "#8B95A1", fontsize=9, weight="normal")
    arrow(ax, (2.95, 1.6), (3.95, 1.6), color="#B0B8C0", lw=1.2)
    arrow(ax, (6.15, 1.6), (7.05, 1.6), color="#B0B8C0", lw=1.2)

    ax.text(
        0.55,
        0.35,
        "本实验中 WMMA GEMM 微基准全部 PASS；当矩阵规模达到 2048 时，Tensor Core 相对标量 CUDA Core 加速约 34.92×。",
        fontsize=8.7,
        color="#52616B",
    )
    return save_figure(fig, "fig4_wmma_schematic")


def get_arrays(data):
    res = data["resnet18_cifar10"]
    batch_sizes = [str(x) for x in res["settings"]["batch_sizes"]]
    xs = np.array([int(x) for x in batch_sizes])
    throughput = {
        method: np.array([res[method]["throughput"][b] for b in batch_sizes], dtype=float)
        for method in METHODS
    }
    accuracy = {method: res[method]["accuracy"] for method in METHODS}
    return xs, throughput, accuracy


def fig_throughput_trends(data):
    xs, throughput, _ = get_arrays(data)
    fig, ax = plt.subplots(figsize=(8.6, 4.8))

    for method in METHODS:
        ax.plot(
            xs,
            throughput[method],
            marker="o",
            lw=2.2,
            ms=5.5,
            color=COLORS[method],
            label=METHOD_LABELS[method],
        )

    best_method = max(METHODS, key=lambda m: throughput[m].max())
    best_idx = int(np.argmax(throughput[best_method]))
    ax.scatter([xs[best_idx]], [throughput[best_method][best_idx]], s=95, color=COLORS["accent"], zorder=5)
    ax.annotate(
        f"峰值 {throughput[best_method][best_idx]:.2f} img/s\n{METHOD_LABELS[best_method]}, BS={xs[best_idx]}",
        xy=(xs[best_idx], throughput[best_method][best_idx]),
        xytext=(70, 467),
        arrowprops=dict(arrowstyle="-|>", color=COLORS["accent"], lw=1.2),
        fontsize=8.5,
        color=COLORS["dark"],
    )

    ax.set_xscale("log", base=2)
    ax.set_xticks(xs)
    ax.get_xaxis().set_major_formatter(mpl.ticker.ScalarFormatter())
    ax.set_xlabel("Batch size")
    ax.set_ylabel("吞吐量 (images/s)")
    ax.set_ylim(245, 492)
    ax.grid(axis="y", alpha=0.22)
    ax.legend(ncol=4, loc="upper center", bbox_to_anchor=(0.5, -0.16), columnspacing=1.6, handlelength=2.2)
    ax.set_title("ResNet-18 / CIFAR-10 推理吞吐量对比", fontsize=11, fontweight="bold", loc="left", pad=10)
    fig.subplots_adjust(bottom=0.22)
    return save_figure(fig, "fig5_throughput_trends")


def fig_speedup_accuracy(data):
    xs, throughput, accuracy = get_arrays(data)
    baseline = throughput["Baseline_FP32"]
    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.2), gridspec_kw={"width_ratios": [1.2, 0.8]})
    ax0, ax1 = axes
    add_panel_label(ax0, "a")
    add_panel_label(ax1, "b")

    for method in METHODS[1:]:
        speedup = throughput[method] / baseline
        ax0.plot(xs, speedup, marker="o", lw=2.2, ms=5, color=COLORS[method], label=METHOD_LABELS[method])
    ax0.axhline(1.0, color="#7B8794", lw=1.1, ls="--")
    ax0.set_xscale("log", base=2)
    ax0.set_xticks(xs)
    ax0.get_xaxis().set_major_formatter(mpl.ticker.ScalarFormatter())
    ax0.set_xlabel("Batch size")
    ax0.set_ylabel("相对 Baseline 加速比 (×)")
    ax0.grid(axis="y", alpha=0.22)
    ax0.legend(loc="upper left")
    ax0.set_title("吞吐量加速比", fontsize=11, fontweight="bold", loc="left")

    labels = [METHOD_SHORT[m] for m in METHODS]
    vals = [accuracy[m] for m in METHODS]
    bars = ax1.bar(labels, vals, color=[COLORS[m] for m in METHODS], width=0.66)
    ax1.axhline(90, color=COLORS["accent"], lw=1.2, ls="--")
    ax1.text(2.65, 90.7, "90% 要求线", color=COLORS["accent"], fontsize=8, ha="center")
    ax1.set_ylim(88, 100)
    ax1.set_ylabel("准确率 (%)")
    ax1.set_title("前 50 张 CIFAR-10 样本", fontsize=11, fontweight="bold", loc="left")
    for bar, val in zip(bars, vals):
        ax1.text(bar.get_x() + bar.get_width() / 2, val + 0.35, f"{val:.0f}%", ha="center", fontsize=8.5)
    ax1.grid(axis="y", alpha=0.16)
    return save_figure(fig, "fig6_speedup_accuracy")


def fig_tensorcore_gemm(data):
    gemm = data["tensor_core_gemm"]["results"]
    sizes = np.array([int(k) for k in gemm.keys()])
    basic = np.array([gemm[str(s)]["basic_ms"] for s in sizes])
    wmma = np.array([gemm[str(s)]["wmma_ms"] for s in sizes])
    speed = np.array([gemm[str(s)]["speedup"] for s in sizes])

    fig, axes = plt.subplots(1, 2, figsize=(9.2, 4.2))
    ax0, ax1 = axes
    add_panel_label(ax0, "a")
    add_panel_label(ax1, "b")

    ax0.plot(sizes, basic, marker="o", lw=2.1, color="#7B8794", label="CUDA Core")
    ax0.plot(sizes, wmma, marker="o", lw=2.1, color=COLORS["Optim_WMMA"], label="Tensor Core WMMA")
    ax0.set_xscale("log", base=2)
    ax0.set_yscale("log")
    ax0.set_xticks(sizes)
    ax0.get_xaxis().set_major_formatter(mpl.ticker.ScalarFormatter())
    ax0.set_xlabel("矩阵尺寸 M=N=K")
    ax0.set_ylabel("执行时间 (ms, log)")
    ax0.grid(axis="both", alpha=0.18)
    ax0.legend(loc="upper left")
    ax0.set_title("FP16 GEMM 执行时间", fontsize=11, fontweight="bold", loc="left")

    bars = ax1.bar([str(s) for s in sizes], speed, color=COLORS["Optim_WMMA"], width=0.65)
    ax1.axhline(1.0, color="#7B8794", lw=1.1, ls="--")
    ax1.set_ylabel("Tensor Core 加速比 (×)")
    ax1.set_xlabel("矩阵尺寸")
    ax1.set_title("Tensor Core 相对 CUDA Core 加速", fontsize=11, fontweight="bold", loc="left")
    ax1.grid(axis="y", alpha=0.18)
    for bar, val in zip(bars, speed):
        ax1.text(bar.get_x() + bar.get_width() / 2, val + 0.8, f"{val:.1f}×", ha="center", fontsize=8.3)
    return save_figure(fig, "fig7_tensorcore_gemm")


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(9)


def style_table(table, header_fill="D9EAF7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def set_doc_defaults(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)

    for name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 11.5)]:
        style = styles[name]
        style.font.name = "黑体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(31, 41, 51)


def add_paragraph(doc, text="", style=None, align=None):
    p = doc.add_paragraph(style=style)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(5)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9)
    run.italic = True
    p.paragraph_format.space_after = Pt(8)


def add_figure(doc, path, caption, width_cm=15.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width_cm))
    add_caption(doc, caption)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(10.5)
    return p


def best_summary(data):
    res = data["resnet18_cifar10"]
    out = {}
    for method in METHODS:
        th = res[method]["throughput"]
        best_bs, best_val = max(th.items(), key=lambda kv: kv[1])
        out[method] = (int(best_bs), float(best_val), float(res[method]["accuracy"]))
    return out


def create_summary_table(doc, data):
    summary = best_summary(data)
    table = doc.add_table(rows=1, cols=5)
    hdr = table.rows[0].cells
    for i, h in enumerate(["算子", "准确率", "最佳 batch size", "最佳吞吐量 (img/s)", "相对 baseline 峰值"]):
        set_cell_text(hdr[i], h, bold=True)
    baseline_best = summary["Baseline_FP32"][1]
    for method in METHODS:
        row = table.add_row().cells
        bs, val, acc = summary[method]
        speedup = val / baseline_best
        values = [METHOD_LABELS[method], f"{acc:.2f}%", bs, f"{val:.2f}", f"{speedup:.2f}×"]
        for i, item in enumerate(values):
            set_cell_text(row[i], item, bold=(method == "Optim_FP16" and i in (0, 3, 4)))
    style_table(table)
    return table


def create_throughput_table(doc, data):
    res = data["resnet18_cifar10"]
    batch_sizes = [str(x) for x in res["settings"]["batch_sizes"]]
    table = doc.add_table(rows=1, cols=1 + len(METHODS))
    hdr = table.rows[0].cells
    headers = ["Batch size"] + [METHOD_SHORT[m] for m in METHODS]
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for bs in batch_sizes:
        row = table.add_row().cells
        set_cell_text(row[0], bs)
        vals = [res[m]["throughput"][bs] for m in METHODS]
        best = max(vals)
        for i, val in enumerate(vals, start=1):
            set_cell_text(row[i], f"{val:.2f}", bold=math.isclose(val, best))
    style_table(table)
    return table


def create_gemm_table(doc, data):
    gemm = data["tensor_core_gemm"]["results"]
    table = doc.add_table(rows=1, cols=5)
    headers = ["矩阵尺寸", "CUDA Core 时间 (ms)", "Tensor Core 时间 (ms)", "加速比", "验证"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
    for size, item in gemm.items():
        row = table.add_row().cells
        values = [f"{size}×{size}", f"{item['basic_ms']:.6g}", f"{item['wmma_ms']:.6g}", f"{item['speedup']:.2f}×", item["verify"]]
        for i, val in enumerate(values):
            set_cell_text(row[i], val, bold=(i in (3, 4) and item["verify"] == "PASS"))
    style_table(table)
    return table


def generate_docx(data, fig_paths):
    DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_doc_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("基于 PyTorch 自定义 CUDA 扩展的 ResNet-18 卷积算子优化")
    run.font.name = "黑体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    run.font.size = Pt(18)
    run.bold = True

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("智能计算系统 Lab6 大作业实验报告")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run("实验平台：NVIDIA GeForce RTX 4060 Laptop GPU | CUDA 12.9 | PyTorch 2.2.2+cu121")
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(9.5)

    doc.add_heading("摘要", level=1)
    add_paragraph(
        doc,
        "卷积算子是卷积神经网络推理阶段的核心计算瓶颈，其访存模式、数据精度和矩阵化程度直接影响模型吞吐量。"
        "本实验以 CIFAR-10 图像分类任务中的 ResNet-18 为应用对象，基于 PyTorch C++/CUDA Extension 实现并比较四类 Conv2d 算子："
        "FP32 直接卷积基线、FP32 implicit-GEMM、FP16 implicit-GEMM 以及 WMMA/Tensor Core 前向卷积。"
        "实验结果表明，在前 50 张 CIFAR-10 测试样本上，Baseline FP32、Optim FP32 和 Optim WMMA 均达到 96.00% 准确率，Optim FP16 达到 94.00% 准确率，均满足不低于 90% 的实验要求。"
        "吞吐量方面，Optim FP16 在 batch size 为 32 时取得 459.99 images/s 的最佳结果，相对 baseline 峰值提升 1.53 倍；FP32 implicit-GEMM 与 WMMA 前向实现分别达到 445.29 images/s 和 393.90 images/s。"
        "此外，独立 FP16 GEMM 微基准表明 Tensor Core 在 2048×2048 矩阵乘上获得 34.92 倍加速且验证全部通过。"
        "这些结果说明，面向卷积推理的算子优化需要同时考虑矩阵化映射、片上存储复用、半精度表示和硬件矩阵单元的实际适配边界。"
    )

    kw = doc.add_paragraph()
    kw.add_run("关键词：").bold = True
    kw.add_run("卷积算子；implicit-GEMM；FP16；Tensor Core；WMMA；PyTorch CUDA Extension；ResNet-18")

    doc.add_heading("1 引言", level=1)
    add_paragraph(
        doc,
        "深度卷积神经网络在图像分类、目标检测和边缘智能等任务中被广泛使用。对于 ResNet-18 这类以卷积层为主体的网络，Conv2d 通常占据推理阶段的大部分计算量。"
        "朴素卷积实现虽然逻辑直接，但其访存访问缺乏规律，输入特征和卷积权重难以在片上存储中充分复用，因而难以发挥 GPU 的并行计算能力。"
    )
    add_paragraph(
        doc,
        "为了提升推理吞吐量，本实验围绕卷积到矩阵乘的等价变换展开优化。FP32 implicit-GEMM 将卷积的输出通道维和输出空间维组织成矩阵乘形式，避免显式 im2col 带来的额外内存开销；"
        "FP16 版本进一步降低数据位宽，以减轻显存带宽压力；WMMA 版本则尝试将 Tensor Core 的矩阵乘累加能力引入卷积前向过程。"
        "图 1 总结了本实验的算子优化路线。"
    )
    add_figure(doc, fig_paths["roadmap"], "图 1 算子优化技术路线。从 FP32 直接卷积出发，逐步引入 implicit-GEMM、FP16 和 WMMA/Tensor Core 前向计算。")

    doc.add_heading("2 方法", level=1)
    doc.add_heading("2.1 系统集成流程", level=2)
    add_paragraph(
        doc,
        "本项目采用三层集成方式：Python 模型层提供 ResNet-18 和自定义 Conv2d wrapper，C++ extension 负责 tensor 检查与参数封装，CUDA kernel 执行具体卷积计算。"
        "该设计使 Python 侧模型结构基本保持不变，同时允许替换底层卷积算子以进行公平对比。"
    )
    add_figure(doc, fig_paths["workflow"], "图 2 PyTorch 自定义 CUDA 卷积算子的调用链。Python 模型经 pybind11 和 C++ 调度层调用 CUDA kernel。")

    doc.add_heading("2.2 直接卷积与 implicit-GEMM 映射", level=2)
    add_paragraph(
        doc,
        "直接卷积中，每个输出元素由输入通道、卷积核高度和卷积核宽度上的三重循环累加得到。该实现便于作为正确性基线，但相邻线程之间对输入和权重的复用较弱。"
        "implicit-GEMM 将权重张量视为 K×CRS 矩阵，将输入张量按运行时地址映射视为 CRS×N·Oh·Ow 矩阵，从而将卷积转化为矩阵乘形式。"
        "与显式 im2col 不同，implicit-GEMM 不实际构造展开矩阵，而是在 kernel 内根据 CRS 索引计算原始 NCHW 地址。"
    )
    add_figure(doc, fig_paths["algorithm"], "图 3 直接卷积与 implicit-GEMM 的算法差异。后者通过隐式地址映射获得矩阵乘结构，同时避免显式 im2col 的额外内存。")

    doc.add_heading("2.3 FP16 与 WMMA/Tensor Core 前向计算", level=2)
    add_paragraph(
        doc,
        "FP16 implicit-GEMM 保持 FP32 版本的分块和共享内存复用结构，仅将输入、权重和中间片段替换为半精度数据表示。"
        "在内存带宽受限的卷积推理中，半精度表示可以降低数据搬运开销，但也可能带来数值舍入误差。"
        "WMMA 版本使用 warp 级矩阵片段和 `mma_sync` 执行 16×16×16 的矩阵乘累加，是本实验中真正调用 Tensor Core 的路径。"
    )
    add_figure(doc, fig_paths["wmma"], "图 4 WMMA/Tensor Core 执行示意。一个 warp 协同加载 A/B 片段并通过 `mma_sync` 完成矩阵乘累加。")

    doc.add_heading("3 实验设置", level=1)
    add_paragraph(
        doc,
        "实验模型为 CIFAR-10 分类任务上的 ResNet-18，输入为 3×32×32 图像，输出为 10 类预测。"
        "准确率测试使用 CIFAR-10 测试集前 50 张图像，吞吐量测试覆盖 batch size 为 8、16、32、64、128、256 和 512 的设置，每个 batch size 重复运行 3 次并取平均吞吐量。"
    )

    env = data["environment"]
    table = doc.add_table(rows=1, cols=2)
    set_cell_text(table.rows[0].cells[0], "项目", bold=True)
    set_cell_text(table.rows[0].cells[1], "配置", bold=True)
    env_rows = [
        ("GPU", env["gpu"]),
        ("Driver CUDA", env["driver_cuda"]),
        ("nvcc", env["nvcc"]),
        ("Python", env["python"]),
        ("PyTorch", env["torch"]),
        ("PyTorch CUDA", env["torch_cuda"]),
        ("准确率样本数", data["resnet18_cifar10"]["settings"]["accuracy_samples"]),
        ("吞吐量重复次数", data["resnet18_cifar10"]["settings"]["run_times"]),
    ]
    for k, v in env_rows:
        row = table.add_row().cells
        set_cell_text(row[0], k)
        set_cell_text(row[1], v)
    style_table(table)
    add_caption(doc, "表 1 实验环境与评测协议。")

    doc.add_heading("4 实验结果与分析", level=1)
    doc.add_heading("4.1 总体性能", level=2)
    add_paragraph(
        doc,
        "表 2 汇总了四类算子的最佳吞吐量与准确率。FP16 implicit-GEMM 获得最高峰值吞吐量，说明在当前模型和输入规模下，降低数据位宽带来的带宽收益较为明显。"
        "FP32 implicit-GEMM 在保持 96.00% 准确率的同时也显著超过 baseline，验证了矩阵化映射和片上存储复用的有效性。"
    )
    create_summary_table(doc, data)
    add_caption(doc, "表 2 四类卷积算子的准确率与最佳吞吐量汇总。")

    add_paragraph(doc, "不同 batch size 下的吞吐量趋势如图 5 和表 3 所示。")
    add_figure(doc, fig_paths["throughput"], "图 5 四类卷积算子在不同 batch size 下的推理吞吐量。FP16 implicit-GEMM 在 batch size 为 32 时达到最高吞吐量。")
    create_throughput_table(doc, data)
    add_caption(doc, "表 3 不同 batch size 下的平均吞吐量，单位为 images/s。")

    doc.add_heading("4.2 加速比与准确率", level=2)
    add_paragraph(
        doc,
        "图 6 同时展示了相对 baseline 的加速比和前 50 张样本上的准确率。所有实现均满足 90% 准确率要求。"
        "Optim FP16 的准确率为 94.00%，比 FP32 与 WMMA 少 1 个正确样本，反映半精度累加和舍入可能对个别边界样本产生影响。"
    )
    add_figure(doc, fig_paths["speedup"], "图 6 相对 baseline 的吞吐量加速比与准确率对比。虚线为实验要求的 90% 准确率基线。")

    doc.add_heading("4.3 Tensor Core GEMM 微基准", level=2)
    add_paragraph(
        doc,
        "为验证 Tensor Core 路径是否正确调用，本实验单独编译运行 `test_mma.cu`，比较标量 CUDA Core FP16 GEMM 与 WMMA GEMM。"
        "结果显示，小规模 128×128 矩阵上 WMMA 受 fragment 加载和 kernel 开销影响不占优；当矩阵规模增大后，Tensor Core 的硬件矩阵乘累加优势迅速显现。"
    )
    add_figure(doc, fig_paths["gemm"], "图 7 Tensor Core 与 CUDA Core 的 FP16 GEMM 微基准。所有尺寸均验证 PASS，2048×2048 时加速比达到 34.92×。")
    create_gemm_table(doc, data)
    add_caption(doc, "表 4 Tensor Core GEMM 微基准结果。")

    doc.add_heading("5 讨论", level=1)
    add_paragraph(
        doc,
        "实验结果表明，implicit-GEMM 对卷积推理的收益主要来自两个方面：一是将卷积映射到更规则的矩阵乘访问模式，二是通过共享内存和寄存器 tile 提高数据复用。"
        "然而，ResNet-18 在 CIFAR-10 上的单层卷积规模较小，kernel launch、边界判断和小矩阵 tile 不饱和等因素会限制峰值吞吐量。"
    )
    add_paragraph(
        doc,
        "WMMA 在独立 GEMM 微基准中表现出非常高的加速比，但在完整卷积前向中未超过 FP16 implicit-GEMM。"
        "这说明 Tensor Core 的理论能力需要足够规整、足够大的矩阵片段才能充分发挥；卷积中的 padding、通道维、输出空间映射和 tile 边界都会引入额外开销。"
        "因此，面向真实卷积层的 Tensor Core 优化不只是替换乘加指令，还需要重新设计数据布局、片段加载和多层级缓存策略。"
    )
    add_paragraph(
        doc,
        "本实验的准确率评估使用前 50 张测试图像，能够验证算子输出在课程要求下的基本可用性，但不足以作为完整模型精度的统计结论。"
        "若进一步提高报告严谨性，应在完整 CIFAR-10 测试集上报告 Top-1 准确率，并加入与 `torch.nn.functional.conv2d` 的逐层数值误差对比。"
    )

    doc.add_heading("6 结论", level=1)
    add_paragraph(
        doc,
        "本文围绕 ResNet-18 推理中的 Conv2d 算子，实现并评测了 FP32 直接卷积、FP32 implicit-GEMM、FP16 implicit-GEMM 和 WMMA/Tensor Core 前向卷积四类 CUDA 算子。"
        "最新实测表明，FP16 implicit-GEMM 获得最高峰值吞吐量 459.99 images/s，FP32 implicit-GEMM 保持 96.00% 准确率并达到 445.29 images/s，WMMA 前向卷积达到 393.90 images/s。"
        "独立 Tensor Core GEMM 微基准进一步证明 WMMA 实现正确，2048×2048 矩阵乘取得 34.92× 加速。"
        "整体来看，卷积算子优化需要在算法映射、内存层次、数据精度和硬件矩阵单元之间进行协同设计。"
    )

    doc.add_heading("参考文献", level=1)
    references = [
        "K. He, X. Zhang, S. Ren, and J. Sun, Deep Residual Learning for Image Recognition, CVPR, 2016.",
        "NVIDIA, CUDA C++ Programming Guide: Warp Matrix Functions and WMMA API.",
        "PyTorch Documentation, C++ and CUDA Extensions.",
        "A. Krizhevsky, Learning Multiple Layers of Features from Tiny Images, 2009.",
    ]
    for ref in references:
        add_number(doc, ref)

    doc.add_heading("附录：复现实验命令", level=1)
    add_paragraph(doc, "构建自定义算子：")
    add_paragraph(doc, "python ./pytorch/setup.py build_ext --inplace")
    add_paragraph(doc, "运行推理 benchmark：")
    add_paragraph(doc, "python inference.py --model fp16 --run-times 3 --accuracy-samples 50 --no-progress")
    add_paragraph(doc, "运行 Tensor Core GEMM 微基准：")
    add_paragraph(doc, "nvcc -arch=sm_89 -O3 -o test_mma.exe test_mma.cu && test_mma.exe")

    doc.save(DOC_PATH)
    return DOC_PATH


def main():
    configure_matplotlib()
    data = load_results()
    fig_paths = {
        "roadmap": fig_optimization_roadmap(),
        "workflow": fig_extension_workflow(),
        "algorithm": fig_algorithm_mapping(),
        "wmma": fig_wmma_schematic(),
        "throughput": fig_throughput_trends(data),
        "speedup": fig_speedup_accuracy(data),
        "gemm": fig_tensorcore_gemm(data),
    }
    doc_path = generate_docx(data, fig_paths)
    print(f"Generated {doc_path}")
    for key, path in fig_paths.items():
        print(f"{key}: {path}")


if __name__ == "__main__":
    main()
